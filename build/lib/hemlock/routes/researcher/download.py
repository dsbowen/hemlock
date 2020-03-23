"""Download

This module performs the following download processes:
1. Create a download page.
2. Create data download files.
3. Create survey view download files.
4. Zip download files.

TODO
Make download survey viewer a class
"""

from hemlock.routes.researcher.utils import *
from hemlock.routes.researcher.login import researcher_login_required
from hemlock.tools import chromedriver

from docx import Document
from docx.shared import Inches

from bs4 import BeautifulSoup
from base64 import b64encode
from datetime import timedelta
from io import BytesIO
from itertools import chain
from zipfile import ZipFile
import os

# Screen buffer for driver window height 
HEIGHT_BUFFER = 74
# Width of survey view for docx
SURVEY_VIEW_IMG_WIDTH = Inches(6)

"""1. Create a download page"""
@bp.route('/download', methods=['GET','POST'])
@researcher_login_required
def download():
    """Download data"""
    return render(download_page())

@researcher_page('download')
def download_page():
    """Return the Download page
    
    The download page allows researchers to select data download files and 
    survey view download files. Data download files are selected as a 
    multi-choice question. Survey view download files are selected by 
    inputing participant IDs in a free response question. IDs are delimited 
    by commas or spaces.
    """
    download_p = Page(navbar=researcher_navbar(), back=False, forward=False)
    
    Check(
        download_p, 
        label=SELECT_FILES_TXT,
        choices=['Metadata', 'Dataframe'],
        multiple=True
    )
    # Choice(data_q, label='Status Log', value='status')

    survey_view_q = Input(download_p, label=SURVEY_VIEW_TXT)
    Validate(survey_view_q, valid_part_ids)
    Submit(survey_view_q, store_part_ids)

    present_q = Select(download_p, label=SURVEY_VIEW_PRESENTATION_TXT)
    Option(present_q, label='First presentation only', value='first')
    Option(present_q, label='All presentations', value='all')

    file_q = Select(download_p, label=SURVEY_VIEW_FILE_TXT)
    Option(file_q, label='Screenshots', value='screenshot')
    Option(file_q, label='Raw text', value='text')

    btn = Download(
        download_p,
        callback=request.url,
        download_msg='Download Complete'
    )
    HandleForm(btn, handle_download_form)
    return download_p

def valid_part_ids(survey_view_q):
    """Check that input participant IDs are valid
    
    Participant IDs must reference participants in the database.
    """
    part_ids = parse_part_ids(survey_view_q.response)
    invalid_ids = [i for i in part_ids if Participant.query.get(i) is None]
    if invalid_ids:
        if len(invalid_ids) == 1:
            return INVALID_ID.format(invalid_ids[0])
        return INVALID_IDS.format(', '.join(invalid_ids))

def store_part_ids(survey_view_q):
    """Store participant IDs in survey view question data"""
    survey_view_q.data = parse_part_ids(survey_view_q.response)

def parse_part_ids(raw_ids):
    """Parse requested participant IDs

    Raw IDs are delimited by spaces or commas. Parsed IDs are a list of 
    integers.
    """
    if not raw_ids:
        return []
    comma_splits = raw_ids.split(',')
    space_splits = [cs.split(' ') for cs in comma_splits]
    return [id for split in space_splits for id in split if id]

def handle_download_form(btn, response):
    """Process download file selection
    
    This method first records the file selection form response and ensures 
    it is valid. If so, it adds the appropriate file creation functions to 
    the download button.
    """
    download_p = download_page()
    db.session.add(download_p)
    btn.downloads = []
    btn.create_file_functions = []
    download_p._record_response()
    data_q, view_q, present_q, file_q, _ = download_p.questions
    if download_p._validate():
        download_p._submit()
        select_data(btn, data_q.data)
        select_survey_view(btn, view_q.data, present_q.data, file_q.data)
        CreateFile(btn, zip_files)
    db.session.commit()

def select_data(btn, files):
    """Add file creation functions for selected data download files"""
    if files.get('Metadata'):
        CreateFile(btn, create_meta)
    if files.get('Status Log'):
        CreateFile(btn, create_status)
    if files.get('Dataframe'):
        CreateFile(btn, create_data)

def select_survey_view(btn, part_ids, present, file_type):
    """Add file creation function for selected survey view download files"""
    if not part_ids:
        return
    CreateFile(
        btn, 
        create_views, 
        args=part_ids, 
        kwargs={'present': present, 'file_type': file_type}
    )

"""2. Create data download files"""
def create_meta(btn):
    """Create metadata download file"""
    stage = 'Preparing Metadata'
    yield btn.reset(stage, 0)
    add_dataframe_to_zip(btn, DataStore.query.first().meta)
    yield btn.report(stage, 100)

def create_status(btn):
    """Create participant status log download file"""
    stage = 'Preparing Status Log'
    yield btn.reset(stage, 0)
    add_dataframe_to_zip(btn, DataStore.query.first().status_log)
    yield btn.report(stage, 100)

def create_data(btn):
    """Create dataframe download file

    Some participants will have updated data. Store data from these 
    participants before creating the dataframe.
    """
    stage = 'Storing Participant Data'
    yield btn.reset(stage, 0)
    ds = DataStore.query.first()
    db.session.add(ds)
    to_store = get_parts_to_store(ds)
    db.session.add_all(to_store)
    for i, part in enumerate(to_store):
        yield btn.report(stage, 100.0*i/len(to_store))
        ds.store_participant(part)
    add_dataframe_to_zip(btn, ds.data)
    db.session.commit()
    yield btn.report(stage, 100)

def get_parts_to_store(ds):
    """Return a list of participants to store in DataStore"""
    all_parts = Participant.query.all()
    parts_stored = ds.parts_stored
    return [
        p for p in all_parts if p.updated or p not in parts_stored
    ]

def add_dataframe_to_zip(btn, data_frame):
    """Add dataframe to zip file creation function"""
    zipfunc = btn.create_file_functions[-1]
    zipfunc.args.append(data_frame.get_download_file())

"""3. Create survey view download files"""
def create_views(btn, *part_ids, present, file_type):
    """Create survey views for all selected participants"""
    parts = [Participant.query.get(id) for id in part_ids]
    driver = chromedriver(headless=True) if file_type=='screenshot' else None
    gen = chain.from_iterable(
        [create_view(btn, p, present, driver) for p in parts]
    )
    for event in gen:
        yield event
    if driver is not None:
        driver.close()

def create_view(btn, part, present, driver):
    """Create survey view for a single participant"""
    stage = 'Creating Survey View for Participant {}'.format(part.id)
    yield btn.reset(stage, 0)
    doc = Document()
    if present == 'all':
        pages = part._viewing_pages
    elif present == 'first':
        pages = [p for p in part._viewing_pages if p.first_presentation]
    for i, page in enumerate(pages):
        yield btn.report(stage, 100.0*i/len(pages))
        add_page_to_doc(doc, page, driver)
    add_doc_to_zip(btn, doc, part.id)
    yield btn.report(stage, 100)

def add_page_to_doc(doc, page, driver):
    """Add survey view page to survey view doc"""
    page.process()
    if driver is None:
        text = BeautifulSoup(page.html, 'html.parser').text.strip('\n')
        doc.add_paragraph(text)
        return
    driver.get('data:text/html,'+page.html)
    accept_alerts(driver)
    width = driver.get_window_size()['width']
    height = driver.execute_script(
        'return document.body.parentNode.scrollHeight'
    )
    driver.set_window_size(width, height+HEIGHT_BUFFER)
    form = driver.find_element_by_tag_name('form')
    page_bytes = BytesIO()
    page_bytes.write(form.screenshot_as_png)
    doc.add_picture(page_bytes, width=SURVEY_VIEW_IMG_WIDTH)
    page_bytes.close()

def accept_alerts(driver):
    """Accept any alerts which prevent page HTML from loading"""
    try:
        while True:
            driver.switch_to.alert.accept()
    except:
        pass

def add_doc_to_zip(btn, doc, part_id):
    """Add survey view doc to download zip file"""
    filename = 'Participant-{}.docx'.format(part_id)
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    zipfunc = btn.create_file_functions[-1]
    zipfunc.args.append((filename, doc_bytes))

"""4. Zip download files"""
def zip_files(btn, *files):
    """Zip download files
    
    `files` is a list of (download_filename, bytes) tuples.
    """
    stage = 'Zipping Files'
    yield btn.reset(stage, 0)
    zipf_bytes = BytesIO()
    zipf = ZipFile(zipf_bytes, 'w')
    for i, (filename, io) in enumerate(files):
        yield btn.report(stage, 100.0*i/len(files))
        zipf.writestr(filename, io.getvalue())
        io.close()
    zipf.close()
    data = b64encode(zipf_bytes.getvalue()).decode()
    url = 'data:application/zip;base64,' + data
    btn.tmp_downloads = [(url, 'download.zip')]
    zipf_bytes.close()
    yield btn.report(stage, 100)