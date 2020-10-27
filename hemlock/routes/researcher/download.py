"""Data download"""

from ...app import bp, db
from ...models import Page, Participant
from ...qpolymorphs import Check, Download, Input, Select
from ...models.private import DataStore
from ...tools import chromedriver, join, show_on_event
from .login import researcher_login_required
from .utils import navbar

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from flask import request, session

import os
import re
import time
from base64 import b64encode
from datetime import timedelta
from io import BytesIO
from itertools import chain
from zipfile import ZipFile

# Screen buffer for driver window height 
HEIGHT_BUFFER = 74
# Width of survey view for docx
SURVEY_VIEW_IMG_WIDTH = Inches(6)
# number of seconds for <iframe> to load before taking screenshot
IFRAME_LOAD_TIME = 3

"""1. Create a download page"""
@bp.route('/download', methods=['GET','POST'])
@researcher_login_required
def download():
    page = get_download_page()
    db.session.commit()
    session['download-page-id'] = page.id
    return page._render()

def get_download_page():
    if 'download-page-id' in session:
        download_page = Page.query.get(session['download-page-id'])
        if download_page:
            return download_page
    
    part_ids_q = Input(
        "Enter participant IDs for survey viewing.",
        validate=valid_part_ids
    )
    presentation_q = Select(
        "Select presentation for viewing.",
        [('First presentation only', 'first'), ('All presentations', 'all')],
    )
    show_on_event(presentation_q, part_ids_q, '\S+', regex=True)
    file_type_q = Select(
        "Select the type of download.",
        [('Screenshots', 'screenshot'), ('Text', 'text')],
    )
    show_on_event(file_type_q, part_ids_q, '\S+', regex=True)

    page = Page(
        Check(
            "Select files to download.", 
            [('Data frame', 'data'), ('Participant metadata', 'meta')],
            multiple=True,
        ),
        part_ids_q,
        presentation_q,
        file_type_q,
        Download(
            callback=request.url, 
            download_msg='Download complete',
            handle_form_functions=handle_download_form,
        ),
        navbar=navbar.render(), back=False, forward=False
    )
    db.session.add(page)
    return page

def valid_part_ids(survey_view_q):
    """
    Validate that requested participant IDs correspond to participants in 
    the database.

    Parameters
    ----------
    survey_view_q : hemlock.Input
        Input question where researchers input participant IDs.

    Returns
    -------
    error_msg : str or None
        A message displaying invalid participant IDs. If this message is 
        not `None`, the download will stop. Otherwise, it proceeds.
    """
    part_ids = re.split(r'\W+', survey_view_q.response)
    part_ids = [id for id in part_ids if id]
    invalid_ids = [i for i in part_ids if Participant.query.get(i) is None]
    if invalid_ids:
        return "<p>The following participant ID was invalid: {}.</p>".format(
            join('and', *invalid_ids)
        )

def handle_download_form(response, btn):
    """
    Adds the appropriate file creation function to the download button.

    Parameters
    ----------
    response : dict
        Form response; not used.

    btn : hemlock.Download
        Download button (last question on the download page).
    """
    page = get_download_page()
    btn.create_file_functions.clear()
    btn.downloads.clear()
    page._record_response()
    if page._validate():
        btn.create_file_functions = FileCreator(
            dataframes=page.questions[0].response,
            part_ids=page.questions[1].response,
            presentation=page.questions[2].response,
            download_type=page.questions[3].response
        )


class FileCreator():
    """
    Creates the download zip file.

    Parameters
    ----------
    dataframes : dict
        Maps dataframe name (`'data'` or `'meta'`) to bool indicating that the
        researcher requested this dataframe.

    part_ids : list
        List of participant IDs requested for survey view.

    presentation : str
        Which page presentation to deliver for the survey view; first 
        presentation only or all presentations.

    download_type : str
        Type of download for the survey view; screenshot or text.

    Attributes
    ----------
    btn : hemlock.Download
    
    datastore : hemlock.models.DataStore

    driver : selenium.webdriver.chrome.webdriver.WebDriver
        Driver for taking screenshots for survey view.

    download_type : str
        Set from parameter.
    
    files : list of (filename, str [bytes]) tuples
        Requested download files.

    presentation : str
        Set from paramaeter.

    requested_dataframes : dict
        Set from `dataframes` parameter.

    requested_part : list of hemlock.Participant
        Participants whose survey view the researcher requested.
    """
    def __init__(self, dataframes, part_ids, presentation, download_type):
        self.btn = None
        self.datastore = None
        self.driver = None
        self.files = []

        self.requested_dataframes = dataframes
        part_ids = re.split(r'\W+', part_ids)
        part_ids = [id for id in part_ids if id]
        self.requested_parts = Participant.query.filter(
            Participant.id.in_(part_ids)
        ).all()
        self.presentation = presentation
        self.download_type = download_type

    def __call__(self, btn):
        """
        Create and deliver the zip file with requested data files.

        Parameters
        ----------
        btn : hemlock.Download
        """
        # note: need to add requested participants before adding `to_store`
        # otherwise you get an error that you try to add the same participant
        # to the session multiple times.
        # I don't know why changing the order in which you add these lists
        # fixes the problem.
        db.session.add_all(self.requested_parts)
        self.btn = btn
        self.datastore = DataStore.query.first()
        if 'data' in self.requested_dataframes:
            for expr in self.prep_dataframe():
                yield expr
        if 'meta' in self.requested_dataframes:
            for expr in self.prep_meta():
                yield expr
        if self.requested_parts:
            for expr in self.gen_views():
                yield expr
        for expr in self.zip_files():
            yield expr

    def prep_dataframe(self):
        """
        Prepare the main data frame for download. Begin by storing data from
        participants whose data was updated since they were last stored in the
        data frame, or whose data has not yet been stored.
        """
        stage = 'Preparing data frame'
        yield self.btn.reset(stage, 0)
        parts = Participant.query.all()
        to_store = [
            p for p in parts 
            if p.updated or p not in self.datastore.parts_stored
        ]
        db.session.add_all(to_store)
        for i, part in enumerate(to_store):
            yield self.btn.report(stage, 100.*i/len(to_store))
            self.datastore.store_participant(part)
        self.files.append(self.datastore.data.get_download_file())
        yield self.btn.report(stage, 100)

    def prep_meta(self):
        """
        Prepare the participant metadata file.
        """
        stage = 'Preparing participant metadata'
        yield self.btn.reset(stage, 0)
        self.files.append(self.datastore.meta.get_download_file())
        yield self.btn.report(stage, 100)

    def gen_views(self):
        """
        Generate survey views for all requested participants.
        """
        if self.download_type == 'screenshot':
            self.driver = chromedriver(headless=True)
        generator = chain.from_iterable(
            (self.gen_view(part) for part in self.requested_parts)
        )
        for event in generator:
            yield event
        if self.driver is not None:
            self.driver.close()
            self.driver = None

    def gen_view(self, part):
        """
        Generate survey view for a single participant.

        Parameters
        ----------
        part : hemlock.Participant
        """
        stage = 'Creating survey view for participant '+str(part.id)
        yield self.btn.reset(stage, 0)
        doc = Document()
        if self.presentation == 'all':
            pages = part._viewing_pages
        elif self.presentation == 'first':
            pages = [p for p in part._viewing_pages if p.first_presentation]
        for i, page in enumerate(pages):
            yield self.btn.report(stage, 100.*i/len(pages))
            self.add_page_to_doc(doc, page)
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        self.files.append(('Participant-{}.docx'.format(part.id), doc_bytes))

    def add_page_to_doc(self, doc, page):
        """
        Add a viewing page to the document.

        Parameters
        ----------
        doc : docx.document.Document
            Survey view document for the current participant.

        page : hemlock.models.private.ViewingPage
            Page to add to the document.
        """
        def write_screenshot_to_doc():
            dist = os.environ.get('WSL_DISTRIBUTION')
            self.driver.get('file://'+('wsl$/'+dist+path if dist else path))
            self.accept_alerts()
            width = self.driver.get_window_size()['width']
            # width = 1024
            height = self.driver.execute_script(
                'return document.body.parentNode.scrollHeight'
            )
            form = self.driver.find_element_by_tag_name('form')
            try:
                _ = form.find_element_by_tag_name('iframe')
                time.sleep(IFRAME_LOAD_TIME)  # give the frame time to load
                height = max(height, 768)
            except:
                pass
            self.driver.set_window_size(width, height+HEIGHT_BUFFER)
            with BytesIO() as page_bytes:
                page_bytes = BytesIO()
                page_bytes.write(form.screenshot_as_png)
                doc.add_picture(page_bytes, width=SURVEY_VIEW_IMG_WIDTH)

        if self.driver is None:
            text = BeautifulSoup(page.html, 'html.parser').text.strip('\n')
            doc.add_paragraph(text)
            return
        path = page.mkstmp()
        try:
            write_screenshot_to_doc()
        except:
            pass
        os.remove(path)

    def accept_alerts(self):
        """
        Accept any alerts which may prevent the driver from loading the page 
        html.
        """
        try:
            while True:
                self.driver.switch_to.alert.accept()
        except:
            pass

    def zip_files(self):
        """
        Put all requested files in a zip archive and store in the download
        button's temporary downloads list.
        """
        stage = 'Zipping files'
        yield self.btn.reset(stage, 0)
        with BytesIO() as zipf_bytes:
            zipf = ZipFile(zipf_bytes, 'w')
            for i, (filename, io) in enumerate(self.files):
                yield self.btn.report(stage, 100.*i/len(self.files))
                zipf.writestr(filename, io.getvalue())
                io.close()
            zipf.close()
            data = b64encode(zipf_bytes.getvalue()).decode()
            url = 'data:application/zip;base64,' + data
            self.btn.tmp_downloads = [(url, 'hemlock-survey-data.zip')]
        yield self.btn.report(stage, 100)